import os
import json
import zipfile
import pandas as pd
from google.cloud import language_v1
from google.oauth2 import service_account
from docx import Document

def get_NLP_data(input_texts):
    os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = r"C:\Users\Hooman Deghani\OneDrive\PC\Desktop\API Keys\natural-language-api-381923-52cfb8b44bd7.json"
    credentials = service_account.Credentials.from_service_account_file(os.environ["GOOGLE_APPLICATION_CREDENTIALS"])
    client = language_v1.LanguageServiceClient(credentials=credentials)

    document = Document()
    report_name = input("Enter the report name: ")

    for text_file in input_texts:
        with open(text_file, 'r', encoding='utf-8') as file:
            text = file.read()

        project_name = input(f"What is the project name for file {text_file}? ")

        document.add_paragraph(f"The project name is {project_name}.")

        document.add_heading("Sentiment Analysis and Category", level=1)
        analyze_sentiment(text, client, document)
        analyze_category(text, client, document)

        document.add_heading("Entities Analysis", level=1)
        analyze_entities(text, client, project_name)

        document.add_heading("Sentiment per Phrase", level=1)
        analyze_sentiment_phrases(text, client, project_name)

    document.save(f"{report_name}.docx")
    print(f"Results written to {report_name}.docx")

def analyze_sentiment(text, client, document):
    document_type = language_v1.Document.Type.PLAIN_TEXT
    language_document = language_v1.Document(content=text, type_=document_type)
    sentiment_response = client.analyze_sentiment(request={'document': language_document})

    overall_sentiment_score = sentiment_response.document_sentiment.score
    overall_sentiment_magnitude = sentiment_response.document_sentiment.magnitude

    paragraph = document.add_paragraph()
    paragraph.add_run("The sentiment score for the project is ").bold = True
    paragraph.add_run(f"{overall_sentiment_score}.")
    paragraph.add_run(" The magnitude of this sentiment score is ").bold = True
    paragraph.add_run(f"{overall_sentiment_magnitude}.")

def analyze_category(text, client, document):
    document_type = language_v1.Document.Type.PLAIN_TEXT
    language_document = language_v1.Document(content=text, type_=document_type)
    category_response = client.classify_text(request={'document': language_document})

    if category_response.categories:
        category = category_response.categories[0].name
        confidence = category_response.categories[0].confidence
        paragraph = document.add_paragraph()
        paragraph.add_run("The category of the project is ").bold = True
        paragraph.add_run(f"{category} ")
        paragraph.add_run("with a confidence of ").bold = True
        paragraph.add_run(f"{confidence}.")
    else:
        document.add_paragraph("No category could be determined for this project.")
    
def analyze_entities(text, client, project_name):
    document_type = language_v1.Document.Type.PLAIN_TEXT
    language_document = language_v1.Document(content=text, type_=document_type)
    entities_response = client.analyze_entities(request={'document': language_document})

    entity_data = [(entity.name, entity.salience) for entity in entities_response.entities]
    entity_df = pd.DataFrame(entity_data, columns=['Entity', 'Salience'])

    entity_df.to_csv(f"entities_{project_name}.csv", index=False)

def analyze_sentiment_phrases(text, client, project_name):
    document_type = language_v1.Document.Type.PLAIN_TEXT
    language_document = language_v1.Document(content=text, type_=document_type)
    sentiment_phrases_response = client.analyze_entity_sentiment(request={'document': language_document})

    sentiment_data = [(entity.name, entity.sentiment.score, entity.sentiment.magnitude) for entity in sentiment_phrases_response.entities]
    sentiment_df = pd.DataFrame(sentiment_data, columns=['Phrase', 'Sentiment Score', 'Magnitude'])

    sentiment_df.to_csv(f"sentiment_{project_name}.csv", index=False)

# Example usage
input_texts = [
    r"C:\Users\Hooman Deghani\Python\Data Analysis\HTML\https___www.investopedia.com_updates_first-time-home-buyer_.txt",
    r"C:\Users\Hooman Deghani\Python\Data Analysis\HTML\https___www.nerdwallet.com_article_mortgages_tips-for-first-time-home-buyers.txt",
    r"C:\Users\Hooman Deghani\Python\Data Analysis\HTML\https___www.rocketmortgage.com_learn_first-time-home-buyer-tips.txt",
    r"C:\Users\Hooman Deghani\Python\Data Analysis\HTML\https___www.squareone.ca_resource-centres_home-buying-selling-moving_buying-a-home-for-the-first-time.txt"
]
get_NLP_data(input_texts)

