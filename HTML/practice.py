import os
import json
import zipfile
import pandas as pd
from google.cloud import language_v1
from google.oauth2 import service_account
from docx import Document

def get_NLP_data(input_texts, SQ1_input):
    os.environ["GOOGLE_APPLICATION_CREDENTIALS"] = r"Path\to\your\Google\Service\Account\JSON\key\file.json"

    credentials = service_account.Credentials.from_service_account_file(os.environ["GOOGLE_APPLICATION_CREDENTIALS"])
    client = language_v1.LanguageServiceClient(credentials=credentials)

    normal_sentiments = []

    # Process normal projects
    for text in input_texts:
        project_name = input(f"What is the project name for file {text}? ")

        overall_sentiment, category_data = analyze_overall_sentiment_and_category(text, client)
        normal_sentiments.append(overall_sentiment.score)

        entity_data = analyze_entities(text, client, project_name)
        sentiment_phrases_data = analyze_sentiment_phrases(text, client, project_name)

        create_csv(entity_data, f"entities_{project_name}")
        create_csv(sentiment_phrases_data, f"sentiment_{project_name}")

        zip_files([f"entities_{project_name}.csv", f"sentiment_{project_name}.csv"], f"{project_name}.zip")

        delete_files([f"entities_{project_name}.csv", f"sentiment_{project_name}.csv"])

        write_sentiment_and_categories_to_docx(project_name, overall_sentiment, category_data, "output.docx")

    # Process special project (SQ1)
    SQ1_overall_sentiment, SQ1_category_data = analyze_overall_sentiment_and_category(SQ1_input, client)
    SQ1_sentiment_score = SQ1_overall_sentiment.score
    SQ1_magnitude = SQ1_overall_sentiment.magnitude
    SQ1_category = SQ1_category_data[0][0]
    SQ1_confidence_score = SQ1_category_data[0][1]

    average_sentiment_score_normals = sum(normal_sentiments) / len(normal_sentiments)
    sentiment_difference = SQ1_sentiment_score - average_sentiment_score_normals
    change_status = "Change required" if abs(sentiment_difference) > 0.1 else "Change not required"

    # Write summary to the docx file
    write_summary_to_docx(SQ1_sentiment_score, average_sentiment_score_normals, SQ1_category, SQ1_confidence_score,
                          sentiment_difference, change_status, "output.docx")

    print("Analysis complete.")


def analyze_sentiment(text, client, document):
    document_type = language_v1.Document.Type.PLAIN_TEXT
    language_document = language_v1.Document(content=text, type_=document_type)
    sentiment_response = client.analyze_sentiment(request={'document': language_document})

    overall_sentiment_score = sentiment_response.document_sentiment.score
    overall_sentiment_magnitude = sentiment_response.document_sentiment.magnitude

    paragraph = document.add_paragraph()
    paragraph.add_run("The sentiment score for the project is ").bold = True
    paragraph.add_run(f"{overall_sentiment_score}.")
    paragraph.add_run(" The magnitude of this sentiment score is ").bold = True
    paragraph.add_run(f"{overall_sentiment_magnitude}.")
    
def analyze_category(text, client, document):
    document_type = language_v1.Document.Type.PLAIN_TEXT
    language_document = language_v1.Document(content=text, type_=document_type)
    category_response = client.classify_text(request={'document': language_document})

    if category_response.categories:
        category = category_response.categories[0].name
        confidence = category_response.categories[0].confidence
        paragraph = document.add_paragraph()
        paragraph.add_run("The category of the project is ").bold = True
        paragraph.add_run(f"{category} ")
        paragraph.add_run("with a confidence of ").bold = True
        paragraph.add_run(f"{confidence}.")
    else:
        document.add_paragraph("No category could be determined for this project.")
    
def analyze_entities(text, client, project_name):
    document_type = language_v1.Document.Type.PLAIN_TEXT
    language_document = language_v1.Document(content=text, type_=document_type)
    entities_response = client.analyze_entities(request={'document': language_document})

    entity_data = [(entity.name, entity.salience) for entity in entities_response.entities]
    entity_df = pd.DataFrame(entity_data, columns=['Entity', 'Salience'])

    entity_df.to_csv(f"entities_{project_name}.csv", index=False)

def analyze_sentiment_phrases(text, client, project_name):
    document_type = language_v1.Document.Type.PLAIN_TEXT
    language_document = language_v1.Document(content=text, type_=document_type)
    sentiment_phrases_response = client.analyze_entity_sentiment(request={'document': language_document})

    sentiment_data = [(entity.name, entity.sentiment.score, entity.sentiment.magnitude) for entity in sentiment_phrases_response.entities]
    sentiment_df = pd.DataFrame(sentiment_data, columns=['Phrase', 'Sentiment Score', 'Magnitude'])

    sentiment_df.to_csv(f"sentiment_{project_name}.csv", index=False)

def write_sentiment_and_categories_to_docx(project_name, overall_sentiment, category_data, docx_file):
    document = Document(docx_file) if os.path.exists(docx_file) else Document()

    document.add_heading(f"The project name is {project_name}.", level=1)

    document.add_heading("Sentiment Analysis and Category", level=2)

    p = document.add_paragraph()
    p.add_run(f"The sentiment score for {project_name} is {overall_sentiment.score}. ").bold = True
    p.add_run(f"The magnitude of this sentiment score is {overall_sentiment.magnitude}. ")

    category_text = ""
    for i, (category, confidence) in enumerate(category_data[:3]):
        category_text += f"The {ordinal(i + 1)} category of {project_name} is {category} with a confidence of {confidence}. "

    p = document.add_paragraph()
    p.add_run(category_text).bold = True

    document.save(docx_file)

def write_summary_to_docx(SQ1_sentiment_score, average_sentiment_score_normals, SQ1_category, SQ1_confidence_score, sentiment_difference, change_status, docx_file):
    document = Document(docx_file) if os.path.exists(docx_file) else Document()

    document.add_heading("Summary", level=1)

    p = document.add_paragraph()
    p.add_run(f"The sentiment score difference of SQ1 and competitors is ").bold = True
    p.add_run(f"{sentiment_difference:.2f}. ").bold = True
    p.add_run(f"{change_status}").bold = True

    p = document.add_paragraph()
    p.add_run(f"SQ1 has a sentiment score of ").bold = True
    p.add_run(f"{SQ1_sentiment_score:.2f}. ").bold = True
    p.add_run(f"The average sentiment score of competitors is ").bold = True
    p.add_run(f"{average_sentiment_score_normals:.2f}. ").bold = True

    p = document.add_paragraph()
    p.add_run(f"SQ1 has a category of ").bold = True
    p.add_run(f"{SQ1_category} ").bold = True
    p.add_run(f"with a confidence score of ").bold = True
    p.add_run(f"{SQ1_confidence_score:.2f}. ").bold = True

    document.save(docx_file)

def ordinal(n):
    suffixes = {1: 'st', 2: 'nd', 3: 'rd'}

    if 10 <= n % 100 <= 20:
        suffix = 'th'
    else:
        suffix = suffixes.get(n % 10, 'th')
    return str(n) + suffix



if __name__ == "__main__":
    input_texts = ["text1.txt", "text2.txt", "text3.txt"]  # Replace these with your own file paths
    SQ1_input = "SQ1_text.txt"  # Replace this with the path to the special project text file
    get_NLP_data(input_texts, SQ1_input)